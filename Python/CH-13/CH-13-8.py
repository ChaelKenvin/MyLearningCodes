import os
import docx

# 写入word文档
os.chdir("D:\\Coding\\Python\\CH-13")

doc = docx.Document()
doc.add_paragraph('Hello world123 !')
# doc.save('hello world123.docx')

#添加段落
paraObj1 = doc.add_paragraph('This is a second paragraph')
paraObj2 = doc.add_paragraph('This is yet another paragraph')
paraObj1.add_run('This text is being added to the second paragraph')

# doc.save('hello world1234.docx')

#添加标题
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

# doc.save('hello world12345.docx')

#添加换行符和换页符
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
doc.paragraphs[0].runs[0].add_break(docx.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')

#添加图像
# doc.add_picture('123.png', width = docx.shared.Inches(1), height = docx.shared.Cm(4))