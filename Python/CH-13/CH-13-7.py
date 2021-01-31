import os
import docx

# 读取word文档
os.chdir("D:\\Coding\\Python\\CH-13")

# doc = docx.Document('demo.docx')
# print(len(doc.paragraphs))

# print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].text)

# print(len(doc.paragraphs[1].runs[0].text))
# print((doc.paragraphs[1].runs[1].text))
# print((doc.paragraphs[1].runs[2].text))
# print((doc.paragraphs[1].runs[3].text))
# print((doc.paragraphs[1].runs[0].text))


def getText(filename):  # 获得完整的文本
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(getText('demo.docx'))
